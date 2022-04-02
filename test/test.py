from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# function align

doc = Document()
doc.add_heading("Hi Niyazi!!", 0)
para = doc.add_paragraph().add_run("Dit is best wel cool Niyazi!!")

doc.add_heading("Nog een stukje.", 1)
para_2 = doc.add_paragraph(
    "Eigenlijk moeten we werken maar ik word hier wel een veel betere IT'er van :)!!!!!"
)

doc.save("Kijk_niyazi.docx")
